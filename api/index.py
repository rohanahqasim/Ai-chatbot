import io
import os
import uuid
import json
from pathlib import Path
from typing import Optional

# Load .env.local when running outside Vercel (e.g. uvicorn local dev).
if not os.environ.get("VERCEL"):
    try:
        from dotenv import load_dotenv

        load_dotenv(Path(__file__).resolve().parent.parent / ".env.local")
    except ImportError:
        pass

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pypdf import PdfReader
import docx

from supabase import create_client, Client

from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_ollama import ChatOllama
from langchain_voyageai import VoyageAIEmbeddings
from langchain_community.vectorstores import SupabaseVectorStore
from langchain_community.tools import DuckDuckGoSearchRun
from langchain_core.tools import tool
from langchain_core.documents import Document
from langchain_core.messages import ToolMessage
from langgraph.prebuilt import create_react_agent


# ---------- App ----------

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # tighten this to your frontend domain once deployed
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.exception_handler(Exception)
async def unhandled_exception_handler(_request, exc: Exception):
    if isinstance(exc, HTTPException):
        raise exc
    return JSONResponse(status_code=500, content={"detail": str(exc)})


# ---------- External services ----------
# All config comes from environment variables set in the Vercel project settings.
# GOOGLE_API_KEY, VOYAGE_API_KEY, SUPABASE_URL, SUPABASE_SERVICE_KEY

supabase: Client = create_client(
    os.environ["SUPABASE_URL"],
    os.environ["SUPABASE_SERVICE_KEY"],
)

if os.environ.get("VERCEL"):
    brain = ChatGoogleGenerativeAI(
        model="gemini-1.5-pro",
        temperature=0.2,
        google_api_key=os.environ["GOOGLE_API_KEY"]
    )
else:
    brain = ChatOllama(
        model="mistral",
        temperature=0.2
    )

engine = VoyageAIEmbeddings(model="voyage-3.5")

vault = SupabaseVectorStore(
    client=supabase,
    embedding=engine,
    table_name="documents",
    query_name="match_documents",
)


# ---------- Tools ----------
# NOTE: run_code (PythonREPL) and save_file/read_file were removed.
# Vercel's filesystem is ephemeral/read-only outside /tmp, and letting an
# agent execute arbitrary Python on a public-facing endpoint is a real
# security risk regardless of platform.

@tool
def search_documents(query: str) -> str:
    """Search previously uploaded documents (stored in the vector database) for
    content relevant to the query. Use this whenever the user asks about a file,
    document, or notebook they uploaded earlier."""
    try:
        results = vault.similarity_search(query, k=4)
    except Exception as exc:
        return f"Document search is temporarily unavailable: {exc}"
    if not results:
        return "No relevant documents found in the vault."
    return "\n\n".join(
        f"[source: {r.metadata.get('source', 'unknown')}]\n{r.page_content}"
        for r in results
    )


tools = [DuckDuckGoSearchRun(), search_documents]

SYSTEM_PROMPT = """
You are a helpful chatbot.

Use search_documents only when the user asks about a file uploaded in a
previous message — not when the current message already includes file context.

If the user asks for Python code, give the code directly. You cannot execute code.

Answer greetings, definitions, and simple conversational questions directly in
natural language without using any tool.
"""

agent = create_react_agent(brain, tools, prompt=SYSTEM_PROMPT)


# ---------- File handling ----------

def extract_content(file_bytes: bytes, filename: str) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(file_bytes))
        return "\n".join([p.extract_text() or "" for p in reader.pages])
    if lower.endswith(".docx"):
        d = docx.Document(io.BytesIO(file_bytes))
        return "\n".join([p.text for p in d.paragraphs])
    if lower.endswith(".doc"):
        raise ValueError(
            f"'{filename}' is a legacy .doc file. Please save it as .docx and upload again."
        )
    if lower.endswith((".txt", ".md", ".csv", ".json")):
        return file_bytes.decode("utf-8", errors="ignore")
    raise ValueError(
        f"Unsupported file type for '{filename}'. Upload PDF, DOCX, TXT, MD, CSV, or JSON."
    )


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 150) -> list[str]:
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start = end - overlap
    return [c for c in chunks if c.strip()]


def store_document(file_bytes: bytes, filename: str) -> tuple[str, str | None]:
    """Extract text and try to embed into Supabase. Returns (text, warning|None)."""
    text = extract_content(file_bytes, filename)
    chunks = chunk_text(text)
    if not chunks:
        return text, "No readable text found in the uploaded file."

    try:
        docs = [Document(page_content=c, metadata={"source": filename}) for c in chunks]
        ids = [str(uuid.uuid4()) for _ in range(len(chunks))]
        vault.add_documents(docs, ids=ids)
    except Exception as exc:
        # Still return extracted text so the chat can answer from the upload context.
        return text, f"Could not save to document vault: {exc}"

    return text, None


# ---------- Session history (Supabase, replaces the in-memory dict) ----------
# The in-memory `sessions = {}` dict from the original code cannot survive on
# Vercel: each request may hit a cold, separate serverless instance.

def load_history(session_id: str) -> list[tuple[str, str]]:
    res = (
        supabase.table("chat_messages")
        .select("role, content")
        .eq("session_id", session_id)
        .order("created_at")
        .execute()
    )
    return [(row["role"], row["content"]) for row in res.data]


def save_message(session_id: str, role: str, content: str) -> None:
    supabase.table("chat_messages").insert(
        {"session_id": session_id, "role": role, "content": content}
    ).execute()


def clear_history(session_id: str) -> None:
    supabase.table("chat_messages").delete().eq("session_id", session_id).execute()


def session_title(content: str) -> str:
    text = (content or "New chat").replace("\n", " ").strip()
    return text[:40] + ("..." if len(text) > 40 else "")


def list_sessions_from_db() -> list[dict]:
    res = (
        supabase.table("chat_messages")
        .select("session_id, role, content, created_at")
        .order("created_at")
        .execute()
    )
    info: dict[str, dict] = {}
    for row in res.data:
        sid = row["session_id"]
        if sid not in info:
            info[sid] = {
                "id": sid,
                "title": "New chat",
                "updated_at": row["created_at"],
            }
        info[sid]["updated_at"] = row["created_at"]
        if row["role"] == "user" and info[sid]["title"] == "New chat":
            info[sid]["title"] = session_title(row["content"])
    return sorted(info.values(), key=lambda s: s["updated_at"], reverse=True)


def message_text(content) -> str:
    """Normalize LangChain/Gemini message content to a plain string."""
    if content is None:
        return ""
    if isinstance(content, str):
        stripped = content.strip()
        if stripped.startswith("[") or stripped.startswith("{"):
            try:
                return message_text(json.loads(content))
            except json.JSONDecodeError:
                pass
        return content
    if isinstance(content, list):
        parts = []
        for block in content:
            if isinstance(block, str):
                parts.append(block)
            elif isinstance(block, dict):
                parts.append(block.get("text", ""))
            else:
                parts.append(str(block))
        return "".join(parts)
    return str(content)


# ---------- Routes ----------

@app.post("/chat")
async def chat_endpoint(
    message: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None),
    session_id: str = Form("default"),
):
    prompt = message or "Analyze."
    steps = []

    if file:
        steps.append(f"Reading uploaded file: {file.filename}")
        file_bytes = await file.read()
        try:
            extracted, vault_warning = store_document(file_bytes, file.filename)
        except ValueError as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc
        if vault_warning:
            steps.append(vault_warning)
        else:
            steps.append(f"Stored '{file.filename}' in the document vault")
        prompt += f"\n\nContext from uploaded file '{file.filename}':\n{extracted[:4000]}"

    history = load_history(session_id)
    history.append(("user", prompt))
    save_message(session_id, "user", prompt)

    response = ""
    try:
        for event in agent.stream({"messages": history}):
            node_data = list(event.values())[0]
            if "messages" not in node_data:
                continue
            msg = node_data["messages"][-1]

            tool_calls = getattr(msg, "tool_calls", None)
            if tool_calls:
                for tc in tool_calls:
                    steps.append(f"Calling tool: {tc.get('name', 'unknown')}")
            elif isinstance(msg, ToolMessage):
                steps.append(f"Got result from: {getattr(msg, 'name', 'tool')}")
            elif getattr(msg, "content", None):
                response = message_text(msg.content)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc

    if not response:
        response = "Sorry, I could not generate a response. Please try again."

    response = message_text(response)
    save_message(session_id, "assistant", response)
    steps.append("Finalizing answer")
    return {"response": response, "steps": steps, "session_id": session_id}


@app.post("/new-session")
async def new_session(session_id: str = Form("default")):
    """Clears history for a session — call this when the user starts a new chat."""
    clear_history(session_id)
    return {"cleared": session_id}


@app.get("/sessions")
async def list_sessions():
    """Lists chat sessions saved in Supabase, newest first."""
    return {"sessions": list_sessions_from_db()}


@app.get("/sessions/{session_id}")
async def get_session_messages(session_id: str):
    """Returns all messages for a session so the UI can reload a past chat."""
    res = (
        supabase.table("chat_messages")
        .select("role, content")
        .eq("session_id", session_id)
        .order("created_at")
        .execute()
    )
    messages = [
        {
            "role": row["role"],
            "text": message_text(row["content"]) if row["role"] == "assistant" else (row["content"] or ""),
        }
        for row in res.data
    ]
    return {"session_id": session_id, "messages": messages}


@app.get("/documents")
async def list_documents():
    """Lists every distinct file that's been embedded into the vault, with chunk counts."""
    res = supabase.table("documents").select("metadata").execute()
    sources: dict[str, int] = {}
    for row in res.data:
        meta = row.get("metadata") or {}
        name = meta.get("source", "unknown")
        sources[name] = sources.get(name, 0) + 1
    return {"documents": [{"name": k, "chunks": v} for k, v in sources.items()]}


@app.delete("/documents/{filename}")
async def delete_document(filename: str):
    """Removes every chunk belonging to a given filename from the vault."""
    res = (
        supabase.table("documents")
        .delete()
        .eq("metadata->>source", filename)
        .execute()
    )
    return {"deleted": filename, "chunks_removed": len(res.data)}


# No uvicorn.run() and no __main__ block — Vercel's Python runtime imports
# the `app` object above directly and serves it. Running uvicorn yourself
# would conflict with that.